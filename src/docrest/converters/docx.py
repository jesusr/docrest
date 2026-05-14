"""docx source/target converters."""
from __future__ import annotations

from pathlib import Path

from docrest.converters.base import ConversionError
from docrest.converters.registry import register


@register("docx", "md")
def docx_to_md(source: Path, target: Path) -> Path:
    try:
        import pypandoc
    except ImportError as exc:
        raise ConversionError("pypandoc required for docx->md") from exc
    try:
        pypandoc.convert_file(str(source), "gfm", outputfile=str(target))
    except OSError as exc:
        raise ConversionError(f"pandoc invocation failed: {exc}") from exc
    return target


@register("docx", "txt")
def docx_to_txt(source: Path, target: Path) -> Path:
    try:
        from docx import Document
    except ImportError as exc:
        raise ConversionError("python-docx required for docx->txt") from exc
    doc = Document(str(source))
    lines = [p.text for p in doc.paragraphs]
    target.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")
    return target


@register("txt", "docx")
def txt_to_docx(source: Path, target: Path) -> Path:
    try:
        from docx import Document
    except ImportError as exc:
        raise ConversionError("python-docx required for txt->docx") from exc
    doc = Document()
    for line in source.read_text(encoding="utf-8").splitlines():
        doc.add_paragraph(line)
    doc.save(str(target))
    return target


@register("docx", "pdf")
def docx_to_pdf(source: Path, target: Path) -> Path:
    md_path = target.with_suffix(".intermediate.md")
    docx_to_md(source, md_path)
    try:
        from docrest.converters.markdown import md_to_pdf

        md_to_pdf(md_path, target)
    finally:
        if md_path.exists():
            md_path.unlink()
    return target
